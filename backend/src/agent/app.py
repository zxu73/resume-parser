# mypy: disable - error - code = "no-untyped-def,misc"
import os
import pathlib
import asyncio
import tempfile
from fastapi import FastAPI, Response, Request, File, UploadFile, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from .agent import evaluate_only_graph, rate_only_graph, optimizer_graph
from .tools import analyze_resume_file
import json
from contextlib import asynccontextmanager
from pydantic import BaseModel, Field, field_validator
import uuid

# Timeouts (seconds) — protect against hung upstream LLM calls
_GRAPH_TIMEOUT = 120.0
_CHAT_TIMEOUT = 30.0

# Store uploaded files on disk so they survive hot-reloads and server restarts
_STORE_DIR = pathlib.Path(tempfile.gettempdir()) / "resume_parser_files"
_STORE_DIR.mkdir(parents=True, exist_ok=True)


def _pdf_path(file_id: str) -> pathlib.Path:
    return _STORE_DIR / f"{file_id}.pdf"


def _doc_path(file_id: str) -> pathlib.Path:
    return _STORE_DIR / f"{file_id}.docx"


def _store_pdf(file_id: str, data: bytes) -> None:
    _pdf_path(file_id).write_bytes(data)


def _store_doc(file_id: str, data: bytes) -> None:
    _doc_path(file_id).write_bytes(data)


def _load_pdf(file_id: str) -> bytes | None:
    p = _pdf_path(file_id)
    return p.read_bytes() if p.exists() else None


def _load_doc(file_id: str) -> bytes | None:
    p = _doc_path(file_id)
    return p.read_bytes() if p.exists() else None

# Pydantic models for request/response
class ResumeEvaluationRequest(BaseModel):
    resume_text: str = Field(max_length=30000)
    job_description: str = Field(max_length=50000)

    @field_validator('resume_text', 'job_description')
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError('must not be empty or whitespace-only')
        return v

class ExperienceItem(BaseModel):
    title: str
    company: str
    duration: str
    description: str
    skills: list[str] = []

class SmartResumeRequest(BaseModel):
    resume_text: str = Field(max_length=30000)
    job_description: str = Field(max_length=50000)
    pool_experiences: list[ExperienceItem] = []  # Optional pool of additional experiences

    @field_validator('resume_text', 'job_description')
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError('must not be empty or whitespace-only')
        return v


class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ChatSuggestion(BaseModel):
    current_text: str
    suggested_text: str
    reason: str


class ChatRequest(BaseModel):
    message: str
    conversation_history: list[ChatMessage] = []
    resume_text: str
    job_description: str


# === JOB SEARCH + AUTO-APPLY MODELS ===

class JobListing(BaseModel):
    job_id: str
    job_title: str
    employer_name: str
    job_city: str
    job_state: str
    job_apply_link: str
    job_description_snippet: str
    job_salary_min: float | None = None
    job_salary_max: float | None = None
    is_greenhouse: bool
    gh_board_token: str | None = None
    gh_job_id: str | None = None


class JobSearchRequest(BaseModel):
    job_description: str = Field(max_length=50000)
    location: str = Field(default="", max_length=200)

    @field_validator('job_description')
    @classmethod
    def job_desc_not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError('must not be empty or whitespace-only')
        return v


class JobSearchResponse(BaseModel):
    success: bool
    jobs: list[JobListing]
    query_used: str


class ApplicantProfile(BaseModel):
    first_name: str = Field(max_length=100)
    last_name: str = Field(max_length=100)
    email: str = Field(max_length=254)
    phone: str = Field(max_length=30)
    linkedin_url: str = Field(default="", max_length=500)


class AutoApplyRequest(BaseModel):
    doc_id: str
    gh_board_token: str
    gh_job_id: str
    applicant: ApplicantProfile
    cover_letter: str = Field(default="", max_length=5000)


class AutoApplyResponse(BaseModel):
    success: bool
    status: str
    message: str
    application_id: str | None = None


import re as _gh_re


def _parse_greenhouse_url(url: str) -> tuple[str, str] | None:
    """Extract (board_token, job_id) from a Greenhouse job URL, or return None."""
    m = _gh_re.search(r'greenhouse\.io/([^/?\s]+)/jobs/(\d+)', url)
    return (m.group(1), m.group(2)) if m else None


@asynccontextmanager
async def lifespan(app: FastAPI):
    yield


# Define the FastAPI app
app = FastAPI(lifespan=lifespan, title="Resume Optimizer API", description="AI-powered resume optimization")


# Add CORS middleware for frontend development.
# allow_credentials must be False when using wildcard origins — browsers reject "*" + credentials.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files (built React frontend) - will be added at the end after all routes


# === RESUME ANALYSIS ENDPOINTS ===

@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload and analyze a resume file (PDF, DOC, TXT).
    Returns analysis results and a unique analysis ID.
    """
    try:
        # Validate file type
        allowed_types = ["application/pdf", "application/msword", 
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "text/plain"]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file.content_type}. Supported types: PDF, DOC, DOCX, TXT"
            )
        
        # Read file content
        content = await file.read()
        
        # Determine file extension
        file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else 'pdf'
        
        # Analyze the resume
        analysis_result = analyze_resume_file(content, file_extension)
        
        if not analysis_result.get("success"):
            raise HTTPException(status_code=500, detail=analysis_result.get("error", "Resume analysis failed"))

        analysis_id = str(uuid.uuid4())

        # Persist file to disk so it survives hot-reloads
        is_pdf = file.content_type == "application/pdf"
        is_docx = file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if is_pdf:
            _store_pdf(analysis_id, content)
        elif is_docx:
            _store_doc(analysis_id, content)

        return {
            "success": True,
            "analysis_id": analysis_id,
            "pdf_id": analysis_id if is_pdf else None,
            "doc_id": analysis_id if is_docx else None,
            "filename": file.filename,
            "file_size": len(content),
            "analysis": analysis_result.get("analysis", ""),
            "extracted_text": analysis_result.get("extracted_text", ""),
            "message": "Resume uploaded and analyzed successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Resume upload failed: {str(e)}")



@app.post("/evaluate-resume")
async def evaluate_resume_directly(request: ResumeEvaluationRequest):
    """Step 1: evaluate the resume and return clarifying questions.

    Does NOT run the rewriter — the client should collect answers to
    clarifying_questions and then POST to /finalize-analysis to get suggestions.
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text cannot be empty")

        if not request.job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")

        try:
            result = await asyncio.wait_for(
                evaluate_only_graph.ainvoke({
                    "resume_text": request.resume_text,
                    "job_description": request.job_description,
                }),
                timeout=_GRAPH_TIMEOUT,
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Resume evaluation timed out")

        evaluation_data = result["evaluation"].model_dump()

        # Deterministic job_match_percentage override.
        matching = evaluation_data.get("matching_skills", [])
        missing = evaluation_data.get("missing_skills", [])
        total_skills = len(matching) + len(missing)
        evaluation_data["job_match_percentage"] = (
            round(len(matching) / total_skills * 100, 1) if total_skills > 0 else 0.0
        )

        clarifying_questions = evaluation_data.pop("clarifying_questions", [])

        return {
            "success": True,
            "structured_evaluation": evaluation_data,
            "clarifying_questions": clarifying_questions,
            "workflow_type": "evaluation_with_questions",
            "message": "Resume evaluated. Answer the clarifying questions to get tailored suggestions.",
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Resume evaluation failed: {str(e)}")


class FinalizeAnalysisRequest(BaseModel):
    resume_text: str = Field(max_length=30000)
    job_description: str = Field(max_length=50000)
    evaluation: dict
    answers: list[dict] = []  # [{"question": str, "answer": str}]

    @field_validator('resume_text', 'job_description')
    @classmethod
    def not_empty(cls, v: str) -> str:
        if not v.strip():
            raise ValueError('must not be empty or whitespace-only')
        return v


@app.post("/finalize-analysis")
async def finalize_analysis(request: FinalizeAnalysisRequest):
    """Step 2: given the evaluation and user answers, produce keyword + STAR suggestions."""
    try:
        try:
            result = await asyncio.wait_for(
                rate_only_graph.ainvoke({
                    "resume_text": request.resume_text,
                    "job_description": request.job_description,
                    "missing_skills": request.evaluation.get("missing_skills", []),
                    "strengths": request.evaluation.get("strengths", []),
                    "weaknesses": request.evaluation.get("weaknesses", []),
                    "answers": request.answers,
                }),
                timeout=_GRAPH_TIMEOUT,
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Finalization timed out")

        # Section invariants (non-empty keywords_added in keyword_suggestions,
        # empty in star_suggestions, no bullet in both) are enforced by
        # RatingResponse.normalize_sections at validation time.
        rating_data = result["rating"].model_dump()

        return {
            "success": True,
            "structured_evaluation": request.evaluation,
            "structured_rating": rating_data,
            "workflow_type": "finalized_with_answers",
            "message": "Suggestions generated with user-provided context.",
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Finalization failed: {str(e)}")

@app.post("/analyze-experience-swaps")
async def analyze_experience_swaps(request: SmartResumeRequest):
    """
    Step 1: Analyze and recommend experience swaps without applying them.
    Returns recommendations for user review.
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text cannot be empty")
        
        if not request.job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        # If no pool experiences provided, fall back to regular evaluation
        if not request.pool_experiences:
            return await evaluate_resume_directly(ResumeEvaluationRequest(
                resume_text=request.resume_text,
                job_description=request.job_description
            ))
        
        try:
            result = await asyncio.wait_for(
                optimizer_graph.ainvoke({
                    "resume_text": request.resume_text,
                    "job_description": request.job_description,
                    "pool_experiences": json.dumps([e.model_dump() for e in request.pool_experiences], indent=2),
                }),
                timeout=_GRAPH_TIMEOUT,
            )
        except asyncio.TimeoutError:
            raise HTTPException(status_code=504, detail="Experience optimization timed out")

        optimization_data = result["optimization"].model_dump()
        
        # Return recommendations for user review (don't apply yet)
        return {
            "success": True,
            "optimization_analysis": optimization_data,
            "workflow_type": "experience_analysis",
            "message": f"Found {optimization_data.get('swaps_made', 0)} recommended swap(s). Review and accept to apply.",
            "requires_user_approval": True
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Smart optimization failed: {str(e)}")

# === EXPERIENCE SWAP ON DOCX ===

class ExperienceSwap(BaseModel):
    resume_experience_title: str
    pool_title: str
    pool_company: str = ""
    pool_duration: str = ""
    pool_description: str = ""

class ApplySwapsRequest(BaseModel):
    doc_id: str
    swaps: list[ExperienceSwap]

@app.post("/apply-swaps-docx")
async def apply_swaps_docx(request: ApplySwapsRequest):
    """Apply accepted experience swaps to the Word document.

    For each swap, finds the section headed by the old experience title
    and replaces it with the pool experience content.  Returns a new
    doc_id pointing to the modified file so the frontend can preview it.
    """
    import re as _re
    from docx import Document as DocxDocument
    import io

    data = _load_doc(request.doc_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document not found")

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def norm(t: str) -> str:
        return _re.sub(r"\s+", " ", t).strip()

    def para_full_text(para) -> str:
        return "".join(
            t.text or "" for t in para._element.iter(f"{{{_W_NS}}}t")
        )

    def write_para(para, new_text: str) -> None:
        all_t = list(para._element.iter(f"{{{_W_NS}}}t"))
        if not all_t:
            return
        all_t[0].text = new_text
        all_t[0].set(_XML_SPACE, "preserve")
        for t in all_t[1:]:
            t.text = ""

    doc = DocxDocument(io.BytesIO(data))
    paras = doc.paragraphs

    for swap in request.swaps:
        old_title_norm = norm(swap.resume_experience_title).lower()

        # Find the paragraph that contains the old experience title
        start_idx = None
        for i, p in enumerate(paras):
            if old_title_norm in norm(para_full_text(p)).lower():
                start_idx = i
                break

        if start_idx is None:
            continue

        # Determine the extent of this experience block:
        # from the title paragraph to the paragraph before the next
        # section/experience heading (heuristic: next paragraph whose
        # font is bold or whose style name contains "Heading", or
        # a blank line followed by another bold/heading paragraph).
        end_idx = start_idx + 1
        for j in range(start_idx + 1, len(paras)):
            text = para_full_text(paras[j]).strip()
            if not text:
                end_idx = j
                continue
            # Check if this paragraph looks like a new heading
            is_bold = any(
                r.bold for r in paras[j].runs if r.text.strip()
            ) if paras[j].runs else False
            style_name = (paras[j].style.name or "").lower()
            is_heading = "heading" in style_name
            if (is_bold or is_heading) and j > start_idx + 1:
                break
            end_idx = j + 1

        # Build replacement lines
        new_lines = [swap.pool_title]
        meta_parts = [p for p in [swap.pool_company, swap.pool_duration] if p]
        if meta_parts:
            new_lines.append(" | ".join(meta_parts))
        if swap.pool_description:
            for bullet in swap.pool_description.split("\n"):
                b = bullet.strip()
                if b:
                    new_lines.append(b)

        # Write replacement: reuse existing paragraphs where possible,
        # clear extras
        for k in range(start_idx, end_idx):
            line_idx = k - start_idx
            if line_idx < len(new_lines):
                write_para(paras[k], new_lines[line_idx])
            else:
                write_para(paras[k], "")

    # Save to a NEW doc_id so the original is preserved
    buf = io.BytesIO()
    doc.save(buf)
    new_bytes = buf.getvalue()

    new_doc_id = f"swapped_{request.doc_id}"
    _store_doc(new_doc_id, new_bytes)

    # Also extract text from the modified doc for the evaluation step
    from .tools import extract_text_from_docx
    modified_text = extract_text_from_docx(new_bytes)

    return {
        "success": True,
        "doc_id": new_doc_id,
        "modified_resume_text": modified_text,
    }


# === FILE VIEWER / DOWNLOAD ENDPOINTS ===

@app.get("/resume-pdf/{pdf_id}")
async def serve_resume_pdf(pdf_id: str):
    """Serve the original uploaded PDF so the frontend can display it."""
    data = _load_pdf(pdf_id)
    if not data:
        raise HTTPException(status_code=404, detail="PDF not found")
    return Response(
        content=data,
        media_type="application/pdf",
        headers={"Cache-Control": "no-store"},
    )


@app.get("/resume-doc/{doc_id}")
async def serve_resume_doc(doc_id: str):
    """Serve the original uploaded Word document so the frontend can render it."""
    data = _load_doc(doc_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document not found")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Cache-Control": "no-store"},
    )


class TextReplacement(BaseModel):
    current_text: str
    suggested_text: str


class ModifyPDFRequest(BaseModel):
    pdf_id: str
    replacements: list[TextReplacement]


def _apply_replacement(page, current_text: str, suggested_text: str) -> bool:
    """
    Find current_text on a PDF page and replace it with suggested_text, preserving
    the original font size and color.

    Strategy:
    1. Use the first 40 chars of the (normalised) current_text as a short anchor for
       page.search_for(), which is reliable for short strings even across ligatures.
    2. Walk page.get_text("dict") spans to collect every span that belongs to the
       full bullet (potentially spanning multiple lines).
    3. Union all their bboxes into one combined rect.
    4. Extract font size + colour from the first matching span.
    5. Apply a single add_redact_annot on the combined rect so the whole multi-line
       bullet is whited out and the new text is drawn in the same size/colour.
    """
    import fitz
    import re

    def norm(t: str) -> str:
        return re.sub(r"\s+", " ", t).strip()

    norm_current = norm(current_text)
    # Strip leading bullet / dash symbols for the anchor search
    anchor_clean = re.sub(r"^[•·\-\*\s]+", "", norm_current)
    anchor = anchor_clean[:40].strip()
    if not anchor:
        return False

    instances = page.search_for(anchor)
    if not instances:
        return False

    anchor_rect = instances[0]

    # Walk all spans to find those belonging to this bullet
    text_dict = page.get_text("dict")
    collected_spans = []
    accumulated = ""
    collecting = False
    font_size = 10.0
    font_color_raw = 0  # integer packed RGB

    for block in text_dict.get("blocks", []):
        if "lines" not in block:
            continue
        for line in block["lines"]:
            for span in line["spans"]:
                span_rect = fitz.Rect(span["bbox"])

                # Start collecting once we hit the anchor span
                if not collecting and span_rect.intersects(anchor_rect):
                    collecting = True
                    font_size = span["size"]
                    font_color_raw = span["color"]

                if collecting:
                    collected_spans.append(span)
                    accumulated += span["text"]

                    # Stop once we have matched enough text
                    if norm_current in norm(accumulated):
                        break
            if collected_spans and norm_current in norm(accumulated):
                break
        if collected_spans and norm_current in norm(accumulated):
            break

    if not collected_spans:
        # Fallback: use just the anchor rect with default font properties
        collected_spans = [{"bbox": anchor_rect, "size": font_size, "color": 0}]

    # Build the union of all collected span bboxes
    combined = fitz.Rect(collected_spans[0]["bbox"])
    for span in collected_spans[1:]:
        combined = combined | fitz.Rect(span["bbox"])

    # Convert packed-int colour to (r, g, b) floats (PyMuPDF stores as 0xRRGGBB int)
    c = font_color_raw
    rgb = ((c >> 16) / 255.0, ((c >> 8) & 0xFF) / 255.0, (c & 0xFF) / 255.0)

    page.add_redact_annot(
        combined,
        text=suggested_text,
        fontsize=round(font_size, 1),
        text_color=rgb,
        align=0,  # left-align
    )
    return True


@app.post("/download-modified-pdf")
async def download_modified_pdf(request: ModifyPDFRequest):
    """Apply approved text replacements to the original PDF and return the modified file."""
    data = _load_pdf(request.pdf_id)
    if not data:
        raise HTTPException(status_code=404, detail="PDF not found")

    import fitz  # pymupdf
    doc = fitz.open(stream=data, filetype="pdf")

    for page in doc:
        changed = False
        for rep in request.replacements:
            if _apply_replacement(page, rep.current_text, rep.suggested_text):
                changed = True
        if changed:
            page.apply_redactions()

    return Response(
        content=doc.tobytes(),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=improved-resume.pdf"},
    )


class ModifyDocxRequest(BaseModel):
    doc_id: str
    replacements: list[TextReplacement]


class ApplySuggestionsRequest(BaseModel):
    doc_id: str
    replacements: list[TextReplacement]


@app.post("/apply-suggestions")
async def apply_suggestions(request: ApplySuggestionsRequest):
    """Apply approved text replacements to the stored Word document and save as a new snapshot.

    Unlike /download-modified-docx, this endpoint does NOT stream a file download.
    It writes the modified document back to disk under a fresh doc_id and returns
    that id so the frontend can reload the live document viewer.
    """
    import re
    from docx import Document as DocxDocument
    import io

    data = _load_doc(request.doc_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document not found")

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def norm(t: str) -> str:
        return re.sub(r"\s+", " ", t).strip()

    _BULLET_RE = re.compile(r"^[\s•·–—\u2013\u2014\u2022\u25aa\u25ba*○▪►◆▸\-]+")

    def strip_bullet(t: str) -> str:
        return _BULLET_RE.sub("", t).strip()

    def para_full_text(para) -> str:
        p_elem = para._element
        return "".join(t.text or "" for t in p_elem.iter(f"{{{_W_NS}}}t"))

    def write_para(para, new_text: str) -> None:
        p_elem = para._element
        all_t = list(p_elem.iter(f"{{{_W_NS}}}t"))
        if not all_t:
            return
        all_t[0].text = new_text
        all_t[0].set(_XML_SPACE, "preserve")
        for t in all_t[1:]:
            t.text = ""

    def is_match(body: str, current: str) -> bool:
        if current in body:
            return True
        anchor = current[:40]
        if len(anchor) >= 20 and anchor in body:
            return True
        return False

    def apply_replacement(paras: list, current: str, suggested: str) -> bool:
        norm_current = strip_bullet(norm(current))
        norm_suggested = strip_bullet(norm(suggested))
        if not norm_current:
            return False
        for i, para in enumerate(paras):
            body = strip_bullet(norm(para_full_text(para)))
            if not body:
                continue
            if not is_match(body, norm_current):
                continue
            write_para(para, norm_suggested)
            j = i + 1
            while j < len(paras):
                cont = strip_bullet(norm(para_full_text(paras[j])))
                if len(cont) >= 8 and cont in norm_current:
                    elem = paras[j]._element
                    elem.getparent().remove(elem)
                    j += 1
                else:
                    break
            return True
        return False

    doc = DocxDocument(io.BytesIO(data))

    for rep in request.replacements:
        if apply_replacement(list(doc.paragraphs), rep.current_text, rep.suggested_text):
            continue
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if apply_replacement(list(cell.paragraphs), rep.current_text, rep.suggested_text):
                        break

    buf = io.BytesIO()
    doc.save(buf)
    new_bytes = buf.getvalue()

    new_doc_id = f"edited_{uuid.uuid4().hex[:8]}"
    _store_doc(new_doc_id, new_bytes)

    return {"success": True, "doc_id": new_doc_id}


@app.post("/download-modified-docx")
async def download_modified_docx(request: ModifyDocxRequest):
    """Apply approved text replacements to the original Word document and return it."""
    import re
    from docx import Document as DocxDocument
    import io

    data = _load_doc(request.doc_id)
    if not data:
        raise HTTPException(status_code=404, detail="Document not found")

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

    def norm(t: str) -> str:
        return re.sub(r"\s+", " ", t).strip()

    _BULLET_RE = re.compile(r"^[\s•·–—\u2013\u2014\u2022\u25aa\u25ba*○▪►◆▸\-]+")

    def strip_bullet(t: str) -> str:
        return _BULLET_RE.sub("", t).strip()

    def para_full_text(para) -> str:
        """Get ALL text in the paragraph, including text inside
        hyperlinks, smart tags, and other nested elements that
        para.runs / para.text might miss."""
        p_elem = para._element
        return "".join(
            t.text or "" for t in p_elem.iter(f"{{{_W_NS}}}t")
        )

    def write_para(para, new_text: str) -> None:
        """Replace ALL text in the paragraph with new_text.

        Finds every <w:t> element (including those nested inside
        <w:hyperlink>, <w:smartTag>, etc.), puts the full new text in
        the first one, and blanks out the rest.  This guarantees no
        leftover text from hidden nested elements.
        """
        p_elem = para._element
        all_t = list(p_elem.iter(f"{{{_W_NS}}}t"))
        if not all_t:
            return
        all_t[0].text = new_text
        all_t[0].set(_XML_SPACE, "preserve")
        for t in all_t[1:]:
            t.text = ""

    def is_match(body: str, current: str) -> bool:
        """Check if a paragraph body matches current_text (exact or anchor)."""
        if current in body:
            return True
        anchor = current[:40]
        if len(anchor) >= 20 and anchor in body:
            return True
        return False

    def apply_replacement(paras: list, current: str, suggested: str) -> bool:
        """Find the paragraph(s) matching current_text, replace with
        suggested_text, and clear any continuation paragraphs.

        A single resume bullet can span multiple Word paragraphs
        (Word wraps long text into continuation <w:p> elements).
        After replacing the first matched paragraph, we scan forward
        and clear subsequent paragraphs whose text appears inside the
        original current_text — these are leftover continuations.
        """
        norm_current = strip_bullet(norm(current))
        norm_suggested = strip_bullet(norm(suggested))

        if not norm_current:
            return False

        for i, para in enumerate(paras):
            body = strip_bullet(norm(para_full_text(para)))
            if not body:
                continue

            if not is_match(body, norm_current):
                continue

            # Found the starting paragraph — replace it
            write_para(para, norm_suggested)

            # Remove continuation paragraphs that are leftover line-wrap fragments.
            # Walk forward and delete (not blank) every paragraph whose text is
            # a substring of the original bullet — leaving no empty <w:p> behind.
            j = i + 1
            while j < len(paras):
                cont = strip_bullet(norm(para_full_text(paras[j])))
                if len(cont) >= 8 and cont in norm_current:
                    elem = paras[j]._element
                    elem.getparent().remove(elem)
                    j += 1
                else:
                    break

            return True

        return False

    doc = DocxDocument(io.BytesIO(data))

    for rep in request.replacements:
        if apply_replacement(list(doc.paragraphs), rep.current_text, rep.suggested_text):
            continue
        # Also search inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if apply_replacement(list(cell.paragraphs), rep.current_text, rep.suggested_text):
                        break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=improved-resume.docx"},
    )


# === JOB SEARCH + AUTO-APPLY ENDPOINTS ===

@app.post("/job-search")
async def job_search(request: JobSearchRequest):
    """Search for matching jobs via JSearch API and detect Greenhouse postings."""
    import httpx
    import litellm as _litellm

    # 1. Extract a concise search query from the job description
    try:
        qr = await _litellm.acompletion(
            model=f"openai/{os.getenv('REASONING_MODEL', 'gpt-4o-mini')}",
            messages=[
                {"role": "system", "content": "Extract the job title and top 2 skills as a short search query (5 words max). Return only the query string, no explanation."},
                {"role": "user", "content": request.job_description[:3000]},
            ],
            max_tokens=50,
        )
        query = (qr.choices[0].message.content or "").strip().strip('"').strip("'")
        if not query:
            raise ValueError("empty query")
    except Exception:
        query = request.job_description[:60].strip()

    if request.location:
        query = f"{query} {request.location}"

    # 2. Call JSearch API
    api_key = os.getenv("JSEARCH_API_KEY", "")
    if not api_key:
        raise HTTPException(status_code=500, detail="JSEARCH_API_KEY not configured")

    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(
                "https://jsearch.p.rapidapi.com/search",
                params={"query": query, "num_pages": "1", "date_posted": "all"},
                headers={
                    "X-RapidAPI-Key": api_key,
                    "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
                },
            )
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Job search timed out")
    except httpx.RequestError as e:
        raise HTTPException(status_code=502, detail=f"Job search unavailable: {e}")

    if resp.status_code == 429:
        raise HTTPException(status_code=429, detail="Job search rate limit reached — try again later")
    if resp.status_code != 200:
        raise HTTPException(status_code=502, detail=f"Job search service error: {resp.status_code}")

    data = resp.json().get("data", [])

    # 3. Build JobListing list with Greenhouse detection
    jobs: list[JobListing] = []
    for item in data:
        apply_link = item.get("job_apply_link") or item.get("job_google_link") or ""
        gh = _parse_greenhouse_url(apply_link)
        jobs.append(JobListing(
            job_id=item.get("job_id", str(uuid.uuid4())),
            job_title=item.get("job_title", ""),
            employer_name=item.get("employer_name", ""),
            job_city=item.get("job_city") or "",
            job_state=item.get("job_state") or "",
            job_apply_link=apply_link,
            job_description_snippet=(item.get("job_description") or "")[:300],
            job_salary_min=item.get("job_min_salary"),
            job_salary_max=item.get("job_max_salary"),
            is_greenhouse=gh is not None,
            gh_board_token=gh[0] if gh else None,
            gh_job_id=gh[1] if gh else None,
        ))

    return JobSearchResponse(success=True, jobs=jobs, query_used=query)


@app.post("/auto-apply")
async def auto_apply(request: AutoApplyRequest):
    """Submit a job application to Greenhouse via their Job Board API."""
    import httpx

    # 1. Load the resume docx from disk
    data = _load_doc(request.doc_id)
    if not data:
        raise HTTPException(
            status_code=404,
            detail="Resume file expired — please re-download your resume and try again",
        )

    # 2. Build multipart form fields + resume file
    form_data = {
        "first_name": request.applicant.first_name,
        "last_name": request.applicant.last_name,
        "email": request.applicant.email,
        "phone": request.applicant.phone,
    }
    if request.applicant.linkedin_url:
        form_data["linkedin_profile"] = request.applicant.linkedin_url
    if request.cover_letter:
        form_data["cover_letter"] = request.cover_letter

    files = {
        "resume": (
            "resume.docx",
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    # 3. POST to Greenhouse Job Board API
    gh_url = f"https://boards-api.greenhouse.io/v1/boards/{request.gh_board_token}/jobs/{request.gh_job_id}"
    api_key = os.getenv("GREENHOUSE_API_KEY", "")

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                gh_url,
                data=form_data,
                files=files,
                auth=(api_key, "") if api_key else None,
            )
    except httpx.TimeoutException:
        return AutoApplyResponse(success=False, status="failed", message="Application submission timed out")
    except httpx.RequestError as e:
        return AutoApplyResponse(success=False, status="failed", message=f"Network error: {e}")

    # 4. Parse Greenhouse response
    try:
        resp_json = resp.json()
    except Exception:
        resp_json = {}

    if resp.status_code == 200:
        app_id = str(resp_json.get("id", ""))
        return AutoApplyResponse(success=True, status="applied", message="Application submitted successfully", application_id=app_id)

    # Non-200: surface the error
    error_detail = resp_json.get("message") or resp_json.get("error") or f"Greenhouse returned {resp.status_code}"
    return AutoApplyResponse(success=False, status="failed", message=error_detail)


# === CHATBOT ENDPOINT ===

@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    """Conversational resume improvement assistant.

    Accepts the user's message and conversation history, then returns a reply
    and optional bullet rewrites that the user can add to their review queue.
    """
    import litellm as _litellm

    system_prompt = (
        "You are BetterCV, an expert resume improvement assistant. "
        "Help the user improve their resume for the given job description.\n\n"
        "When the user asks to improve or rewrite specific bullets, identify the relevant "
        "bullet(s) from the resume and suggest improved versions in STAR format:\n"
        "  [Strong past-tense verb] + [technical context/approach] + [result or impact]\n\n"
        "RESUME TEXT:\n"
        f"{request.resume_text}\n\n"
        "JOB DESCRIPTION:\n"
        f"{request.job_description}\n\n"
        "RESPONSE FORMAT — always return valid JSON with exactly these keys:\n"
        '{{\n'
        '  "reply": "Your conversational response to the user",\n'
        '  "suggestions": [\n'
        '    {{\n'
        '      "current_text": "Exact bullet text copied from the resume",\n'
        '      "suggested_text": "Improved STAR-format version",\n'
        '      "reason": "One sentence on why this is better"\n'
        '    }}\n'
        '  ]\n'
        '}}\n\n'
        "Rules:\n"
        "- suggestions must be [] when no specific rewrites are needed (e.g. questions, general discussion)\n"
        "- current_text must be copied verbatim from the resume — do not invent new bullets\n"
        "- Do NOT fabricate metrics, project names, or skills the user never mentioned\n"
        "- Be specific and actionable in your reply\n"
        "- Return ONLY valid JSON — no markdown fences, no extra text"
    )

    messages: list[dict] = [{"role": "system", "content": system_prompt}]
    for msg in request.conversation_history:
        messages.append({"role": msg.role, "content": msg.content})
    messages.append({"role": "user", "content": request.message})

    try:
        response = await _litellm.acompletion(
            model=f"openai/{os.getenv('REASONING_MODEL', 'gpt-4o-mini')}",
            messages=messages,
            response_format={"type": "json_object"},
            timeout=_CHAT_TIMEOUT,
        )
    except Exception as e:
        raise HTTPException(status_code=504, detail=f"Chat request failed: {e}")

    content = response.choices[0].message.content or "{}"
    try:
        data = json.loads(content)
    except json.JSONDecodeError:
        data = {"reply": content, "suggestions": []}

    return {
        "reply": data.get("reply", ""),
        "suggestions": data.get("suggestions", []),
    }


# === FRONTEND STATIC FILES ===
# Mount React frontend after all API routes (must be last)

from pathlib import Path

frontend_dist_path = Path(__file__).parent.parent.parent.parent / "frontend" / "dist"

if frontend_dist_path.exists() and (frontend_dist_path / "index.html").exists():
    from fastapi.responses import HTMLResponse

    from fastapi.responses import FileResponse

    # Serve hashed assets (js/css/images) — these are safe to cache
    app.mount("/assets", StaticFiles(directory=frontend_dist_path / "assets"), name="assets")

    # Serve other static files from dist root (e.g. vite.svg, favicon.ico)
    @app.get("/{file_name:path}")
    async def serve_spa(file_name: str):
        # If the file exists in dist, serve it directly
        file_path = frontend_dist_path / file_name
        if file_name and file_path.is_file() and file_path.name != "index.html":
            return FileResponse(file_path)
        # Otherwise serve index.html with no-cache for SPA routing
        return HTMLResponse(
            content=(frontend_dist_path / "index.html").read_text(),
            headers={
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0",
            },
        )
    print(f"[OK] Frontend mounted from: {frontend_dist_path}")
else:
    print(f"[WARN] Frontend build not found at: {frontend_dist_path}")
    
    @app.get("/")
    async def root():
        return {"message": "Resume Analyzer API", "status": "Backend running", "note": "Frontend build not found"}